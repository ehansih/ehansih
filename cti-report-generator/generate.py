"""
generate.py — CTI Daily Digest Report Generator
Reads manifest.yaml and produces a formatted Word document.
Usage: python3 generate.py [--manifest manifest.yaml] [--out output/report.docx]
"""
import yaml, os, argparse
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import images as img_gen

# ── Colors ────────────────────────────────────────────────────────────────────
C = dict(
    DARK_NAVY='0D1B2A', NAVY='1B3A5C', ACCENT_BLUE='1F6FEB',
    CRITICAL='C0392B', HIGH='E67E22', MEDIUM='F1C40F',
    LOW='27AE60', INFO='2980B9', LIGHT_GRAY='F4F6F8',
    MID_GRAY='D5D8DC', WHITE='FFFFFF', TEXT_DARK='1A1A2E',
    SECTION_HDR='0A3D62',
)

def rgb(h):
    r,g,b = tuple(int(h[i:i+2],16) for i in (0,2,4))
    return RGBColor(r,g,b)

def sev_color(level):
    return {'CRITICAL':C['CRITICAL'],'HIGH':C['HIGH'],
            'MEDIUM':C['MEDIUM'],'LOW':C['LOW']}.get(str(level).upper(), C['INFO'])

# ── Word helpers ──────────────────────────────────────────────────────────────
def set_bg(cell, color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),'clear'); shd.set(qn('w:color'),'auto'); shd.set(qn('w:fill'),color)
    tcPr.append(shd)

def set_borders(cell, color='C8D6E5', size='4'):
    tcPr = cell._tc.get_or_add_tcPr()
    tcB = OxmlElement('w:tcBorders')
    for side in ('top','left','bottom','right'):
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'),'single'); b.set(qn('w:sz'),size)
        b.set(qn('w:space'),'0'); b.set(qn('w:color'),color)
        tcB.append(b)
    tcPr.append(tcB)

def shd_run(run, fill):
    rPr = run._r.get_or_add_rPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),'clear'); shd.set(qn('w:color'),'auto'); shd.set(qn('w:fill'),fill)
    rPr.append(shd)

def add_image(doc, path, width=6.5):
    if not path or not os.path.exists(path):
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(8)
    p.add_run().add_picture(path, width=Inches(width))

def add_table(doc, headers, rows, hbg=None):
    hbg = hbg or C['DARK_NAVY']
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = 'Table Grid'; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i,h in enumerate(headers):
        cell = t.rows[0].cells[i]; set_bg(cell,hbg); set_borders(cell,C['ACCENT_BLUE'],'6')
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h); run.bold=True; run.font.color.rgb=rgb(C['WHITE']); run.font.size=Pt(8)
    for ri,row in enumerate(rows):
        bg = C['LIGHT_GRAY'] if ri%2==0 else C['WHITE']
        for ci,val in enumerate(row):
            cell=t.rows[ri+1].cells[ci]; set_bg(cell,bg); set_borders(cell,C['MID_GRAY'],'4')
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p=cell.paragraphs[0]; run=p.add_run(str(val)); run.font.size=Pt(8)
            if ci==0: run.bold=True
            if ci==3:
                sc={'CRITICAL':C['CRITICAL'],'HIGH':C['HIGH'],'MEDIUM':C['MEDIUM'],'LOW':C['LOW']}
                run.bold=True; run.font.color.rgb=rgb(sc.get(str(val),C['TEXT_DARK']))
            if ci==4:
                rc={'HIGH':C['CRITICAL'],'MEDIUM':C['HIGH'],'LOW':C['LOW']}
                run.bold=True; run.font.color.rgb=rgb(rc.get(str(val),C['TEXT_DARK']))
    doc.add_paragraph().paragraph_format.space_after=Pt(8)

def sub_heading(doc, text):
    p=doc.add_paragraph(); r=p.add_run(text)
    r.bold=True; r.font.size=Pt(11); r.font.color.rgb=rgb(C['SECTION_HDR'])
    p.paragraph_format.space_after=Pt(4)

def add_threat_card(doc, threat):
    severity = str(threat.get('severity','MEDIUM')).upper()
    telecom  = str(threat.get('telecom_relevance','MEDIUM')).upper()

    outer = doc.add_table(rows=1, cols=1); outer.style='Table Grid'
    oc = outer.rows[0].cells[0]
    set_borders(oc, sev_color(severity), '6'); set_bg(oc,'F8FAFC')

    # Title
    tp = oc.paragraphs[0]; tp.paragraph_format.space_before=Pt(4); tp.paragraph_format.space_after=Pt(2)
    tr = tp.add_run(threat.get('title','Untitled')); tr.bold=True; tr.font.size=Pt(11); tr.font.color.rgb=rgb(C['DARK_NAVY'])

    # Metadata row
    meta = oc.add_paragraph(); meta.paragraph_format.space_before=Pt(2); meta.paragraph_format.space_after=Pt(4)
    sr = meta.add_run(f'  {severity}  '); sr.bold=True; sr.font.size=Pt(8); sr.font.color.rgb=rgb(C['WHITE'])
    shd_run(sr, sev_color(severity))
    meta.add_run('   ')
    dr = meta.add_run(f'Date: {threat.get("date","")}'); dr.font.size=Pt(9); dr.font.color.rgb=rgb('5D6D7E')
    meta.add_run('   |   ')
    rr = meta.add_run(f'Telecom Relevance: {telecom}'); rr.bold=True; rr.font.size=Pt(9)
    rr.font.color.rgb=rgb({'HIGH':C['CRITICAL'],'MEDIUM':C['HIGH'],'LOW':C['LOW']}.get(telecom,C['INFO']))

    # CVE / CVSS
    cve_list = threat.get('cve',[])
    cvss_val = threat.get('cvss')
    if cve_list or cvss_val:
        cp = oc.add_paragraph(); cp.paragraph_format.space_before=Pt(0); cp.paragraph_format.space_after=Pt(4)
        if cve_list:
            for c in (cve_list if isinstance(cve_list,list) else [cve_list]):
                cr=cp.add_run(f' {c} '); cr.bold=True; cr.font.size=Pt(8); cr.font.color.rgb=rgb(C['WHITE'])
                shd_run(cr,C['NAVY']); cp.add_run(' ')
        if cvss_val:
            cp.add_run('  ')
            try: cv=float(str(cvss_val).split()[0])
            except: cv=0
            cc=C['CRITICAL'] if cv>=9 else C['HIGH'] if cv>=7 else C['MEDIUM'] if cv>=4 else C['LOW']
            cr2=cp.add_run(f' CVSS: {cvss_val} '); cr2.bold=True; cr2.font.size=Pt(8); cr2.font.color.rgb=rgb(C['WHITE'])
            shd_run(cr2,cc)

    # Body sections
    fields = [
        ('DESCRIPTION',         threat.get('description'),   'text'),
        ('AFFECTED PRODUCTS',   threat.get('affected'),      'text'),
        ('MITRE ATT&CK TTPs',   threat.get('mitre_ttps'),    'ttps'),
        ('RECOMMENDED ACTIONS', threat.get('recommendations'),'recs'),
    ]
    for label,content,mode in fields:
        if not content: continue
        lp=oc.add_paragraph(); lp.paragraph_format.space_before=Pt(4); lp.paragraph_format.space_after=Pt(1)
        lr=lp.add_run(label); lr.bold=True; lr.font.size=Pt(8)
        lr.font.color.rgb=rgb(C['LOW'] if label=='RECOMMENDED ACTIONS' else C['NAVY'])
        if mode=='ttps':
            tp2=oc.add_paragraph(); tp2.paragraph_format.space_after=Pt(6)
            for ttp in content:
                tr2=tp2.add_run(f' {ttp} '); tr2.font.size=Pt(8); tr2.font.color.rgb=rgb(C['WHITE'])
                shd_run(tr2,'4A235A'); tp2.add_run('  ')
        elif mode=='recs':
            for i,rec in enumerate(content,1):
                rp=oc.add_paragraph(); rp.paragraph_format.space_before=Pt(1); rp.paragraph_format.space_after=Pt(1)
                nr=rp.add_run(f'{i}. '); nr.bold=True; nr.font.size=Pt(9); nr.font.color.rgb=rgb(C['LOW'])
                rp.add_run(str(rec)).font.size=Pt(9)
        else:
            cp2=oc.add_paragraph(); cp2.paragraph_format.space_after=Pt(6)
            cp2.add_run(str(content)).font.size=Pt(9)

    if threat.get('reference'):
        rp2=oc.add_paragraph(); rp2.paragraph_format.space_before=Pt(4); rp2.paragraph_format.space_after=Pt(4)
        rr2=rp2.add_run('Reference: '); rr2.bold=True; rr2.font.size=Pt(8); rr2.font.color.rgb=rgb('7F8C8D')
        rp2.add_run(str(threat['reference'])).font.size=Pt(8)

    doc.add_paragraph().paragraph_format.space_after=Pt(8)

# ── Main builder ──────────────────────────────────────────────────────────────
def build_report(manifest_path, output_path, img_dir=None):
    with open(manifest_path,'r') as f:
        m = yaml.safe_load(f)

    report_meta = m.get('report',{})
    org    = report_meta.get('organization','Nokia Networks')
    date   = report_meta.get('date','')
    title  = report_meta.get('title','Threat Intelligence Report — Daily Digest')
    team   = report_meta.get('team','Security Intelligence Center')
    classif= report_meta.get('classification','TLP: AMBER')
    footer = report_meta.get('footer_note','')

    # Generate images
    if img_dir is None:
        img_dir = os.path.join(os.path.dirname(output_path), 'images')
    print(f'Generating section images in {img_dir}...')
    img_paths = img_gen.generate_all(img_dir, org, date)

    doc = Document()
    for section in doc.sections:
        section.top_margin=Cm(1.8); section.bottom_margin=Cm(1.8)
        section.left_margin=Cm(2.0); section.right_margin=Cm(2.0)

    # ── Cover ──────────────────────────────────────────────────────────────────
    add_image(doc, img_paths.get('cover'), 6.5)
    tp=doc.add_paragraph(); tp.alignment=WD_ALIGN_PARAGRAPH.CENTER
    t1=tp.add_run(title); t1.bold=True; t1.font.size=Pt(13); t1.font.color.rgb=rgb(C['SECTION_HDR'])
    op=doc.add_paragraph(); op.alignment=WD_ALIGN_PARAGRAPH.CENTER
    o1=op.add_run(f'{org}  |  {team}  |  {date}')
    o1.font.size=Pt(10); o1.font.color.rgb=rgb('5D6D7E'); o1.italic=True
    tlp_p=doc.add_paragraph(); tlp_p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    tlp_p.paragraph_format.space_before=Pt(6)
    tlp_run=tlp_p.add_run(f'  {classif} — RESTRICTED TO {org.upper()} SOC & AUTHORIZED PERSONNEL  ')
    tlp_run.bold=True; tlp_run.font.size=Pt(9); tlp_run.font.color.rgb=rgb(C['WHITE'])
    shd_run(tlp_run, C['HIGH'])
    doc.add_page_break()

    # ── Executive Summary ──────────────────────────────────────────────────────
    ep=doc.add_paragraph(); er=ep.add_run('EXECUTIVE SUMMARY')
    er.bold=True; er.font.size=Pt(16); er.font.color.rgb=rgb(C['SECTION_HDR'])
    ep.paragraph_format.space_after=Pt(6)

    # Count threats across sections
    all_threats = (
        m.get('vulnerabilities',{}).get('threats',[]) +
        m.get('malware',{}).get('threats',[]) +
        m.get('campaigns',{}).get('threats',[]) +
        m.get('breaches',{}).get('threats',[])
    )
    n_critical = sum(1 for t in all_threats if str(t.get('severity','')).upper()=='CRITICAL')
    n_high     = sum(1 for t in all_threats if str(t.get('severity','')).upper()=='HIGH')
    n_medium   = sum(1 for t in all_threats if str(t.get('severity','')).upper()=='MEDIUM')
    n_telecom  = sum(1 for t in all_threats if str(t.get('telecom_relevance','')).upper() in ('HIGH','MEDIUM'))

    stat_table=doc.add_table(rows=2,cols=5); stat_table.style='Table Grid'
    for i,(h,v,c) in enumerate(zip(
        ['Total Threats','Critical','High','Medium','Telecom-Relevant'],
        [str(len(all_threats)),str(n_critical),str(n_high),str(n_medium),str(n_telecom)],
        [C['NAVY'],C['CRITICAL'],C['HIGH'],C['MEDIUM'],C['SECTION_HDR']]
    )):
        hcell=stat_table.rows[0].cells[i]; vcell=stat_table.rows[1].cells[i]
        set_bg(hcell,c); set_bg(vcell,C['LIGHT_GRAY'])
        set_borders(hcell,C['WHITE'],'6'); set_borders(vcell,C['MID_GRAY'],'4')
        hcell.vertical_alignment=vcell.vertical_alignment=WD_ALIGN_VERTICAL.CENTER
        hp=hcell.paragraphs[0]; hp.alignment=WD_ALIGN_PARAGRAPH.CENTER
        hr=hp.add_run(h); hr.bold=True; hr.font.size=Pt(8); hr.font.color.rgb=rgb(C['WHITE'])
        vp=vcell.paragraphs[0]; vp.alignment=WD_ALIGN_PARAGRAPH.CENTER
        vr=vp.add_run(v); vr.bold=True; vr.font.size=Pt(18); vr.font.color.rgb=rgb(c)
    doc.add_paragraph().paragraph_format.space_after=Pt(12)

    # Overview table
    sub_heading(doc,'THREAT LANDSCAPE OVERVIEW')
    overview_rows=[]
    for i,t in enumerate(all_threats,1):
        cat='Vulnerability' if t in m.get('vulnerabilities',{}).get('threats',[]) else \
            'Malware' if t in m.get('malware',{}).get('threats',[]) else \
            'Campaign' if t in m.get('campaigns',{}).get('threats',[]) else 'Breach'
        overview_rows.append([str(i), t.get('title','')[:60], cat,
                               str(t.get('severity','MEDIUM')).upper(),
                               str(t.get('telecom_relevance','MEDIUM')).upper(),
                               (t.get('recommendations') or ['Review IOCs'])[0][:55]])
    add_table(doc,['#','Threat','Category','Severity','Telecom Relevance','Action Required'],overview_rows)
    doc.add_page_break()

    # ── Section 1: Vulnerabilities ─────────────────────────────────────────────
    add_image(doc, img_paths.get('vuln'), 6.5)
    env_note = m.get('vulnerabilities',{}).get('environment_note','')
    if env_note:
        nt=doc.add_table(rows=1,cols=1); nt.style='Table Grid'
        nc=nt.rows[0].cells[0]; set_bg(nc,'EAF4FB'); set_borders(nc,C['ACCENT_BLUE'],'6')
        np2=nc.paragraphs[0]; np2.paragraph_format.space_before=Pt(4); np2.paragraph_format.space_after=Pt(4)
        n1=np2.add_run(f'{org} Environment Status:  '); n1.bold=True; n1.font.size=Pt(9); n1.font.color.rgb=rgb(C['NAVY'])
        np2.add_run(str(env_note).strip()).font.size=Pt(9)
        doc.add_paragraph().paragraph_format.space_after=Pt(6)
    for t in m.get('vulnerabilities',{}).get('threats',[]):
        add_threat_card(doc,t)
    doc.add_page_break()

    # ── Section 2: Malware ─────────────────────────────────────────────────────
    add_image(doc, img_paths.get('malware'), 6.5)
    for t in m.get('malware',{}).get('threats',[]):
        add_threat_card(doc,t)
    doc.add_page_break()

    # ── Section 3: Campaigns ───────────────────────────────────────────────────
    add_image(doc, img_paths.get('campaigns'), 6.5)
    for t in m.get('campaigns',{}).get('threats',[]):
        add_threat_card(doc,t)
    doc.add_page_break()

    # ── Section 4: Breaches ────────────────────────────────────────────────────
    add_image(doc, img_paths.get('breaches'), 6.5)
    for t in m.get('breaches',{}).get('threats',[]):
        add_threat_card(doc,t)
    doc.add_page_break()

    # ── Section 5: IOCs ────────────────────────────────────────────────────────
    add_image(doc, img_paths.get('ioc'), 6.5)
    ioc_data = m.get('iocs',{})
    ioc_note=doc.add_paragraph(
        f'All IOCs below are derived from threats in this digest. Feed into SIEM, firewall blocklists, '
        f'DNS RPZ, EDR, and email gateway. {classif} — {org} SOC and authorized personnel only. '
        f'Defang notation used ([.]) — reformat before use in blocking tools.')
    ioc_note.runs[0].font.size=Pt(9); ioc_note.runs[0].font.color.rgb=rgb('5D6D7E')
    ioc_note.paragraph_format.space_after=Pt(8)

    if ioc_data.get('domains') or ioc_data.get('ips'):
        sub_heading(doc,'Network IOCs — Domains & IPs')
        rows=[]
        for d in ioc_data.get('domains',[]): rows.append(['Domain',d['indicator'],d['threat'],d['action']])
        for ip in ioc_data.get('ips',[]): rows.append(['IP',ip['indicator'],ip['threat'],ip['action']])
        add_table(doc,['Type','Indicator','Associated Threat','Action'],rows)

    if ioc_data.get('hashes'):
        sub_heading(doc,'File Hash IOCs')
        rows=[[h['type'],h['hash'],h['family'],h['action']] for h in ioc_data['hashes']]
        add_table(doc,['Type','Hash','File / Malware Family','Action'],rows)

    if ioc_data.get('urls'):
        sub_heading(doc,'Malicious URL IOCs')
        rows=[[u['indicator'],'URL',u['threat'],u['action']] for u in ioc_data['urls']]
        add_table(doc,['URL','Type','Associated Threat','Action'],rows)
    doc.add_page_break()

    # ── Section 6: Threat Actors ───────────────────────────────────────────────
    add_image(doc, img_paths.get('actors'), 6.5)
    actors = m.get('threat_actors',[])
    if actors:
        rows=[[a['name'],a['origin'],a['motivation'],a['targets'],a['ttps'],a['activity_level']] for a in actors]
        add_table(doc,['Actor','Origin','Motivation','Primary Targets','TTPs','Activity Level'],rows)
    doc.add_page_break()

    # ── Section 7: Action Items ────────────────────────────────────────────────
    add_image(doc, img_paths.get('actions'), 6.5)
    actions = m.get('action_items',{})
    color_map={'priority_1':C['CRITICAL'],'priority_2':C['HIGH'],'priority_3':C['MEDIUM']}
    for pkey in ('priority_1','priority_2','priority_3'):
        pdata = actions.get(pkey,{})
        if not pdata: continue
        pp=doc.add_paragraph(); pr=pp.add_run(pdata.get('label',''))
        pr.bold=True; pr.font.size=Pt(11); pr.font.color.rgb=rgb(color_map[pkey])
        pp.paragraph_format.space_after=Pt(4)
        for item in pdata.get('items',[]):
            bp=doc.add_paragraph(style='List Bullet'); bp.paragraph_format.space_after=Pt(3)
            r1=bp.add_run(f'{item["title"]}: '); r1.bold=True; r1.font.size=Pt(9); r1.font.color.rgb=rgb(color_map[pkey])
            bp.add_run(str(item['detail'])).font.size=Pt(9)
        doc.add_paragraph().paragraph_format.space_after=Pt(6)

    # ── Footer ─────────────────────────────────────────────────────────────────
    ft=doc.add_table(rows=1,cols=1); ft.style='Table Grid'
    fc=ft.rows[0].cells[0]; set_bg(fc,C['DARK_NAVY']); set_borders(fc,C['ACCENT_BLUE'],'4')
    fp=fc.paragraphs[0]; fp.alignment=WD_ALIGN_PARAGRAPH.CENTER
    fp.paragraph_format.space_before=Pt(6); fp.paragraph_format.space_after=Pt(6)
    fr=fp.add_run(f'{org}  |  {title}  |  {date}  |  {classif}\n{footer}')
    fr.font.size=Pt(8); fr.font.color.rgb=rgb('A9CCE3'); fr.italic=True

    os.makedirs(os.path.dirname(os.path.abspath(output_path)), exist_ok=True)
    doc.save(output_path)
    print(f'Report saved: {output_path}')
    return output_path

if __name__=='__main__':
    parser = argparse.ArgumentParser(description='Generate CTI Daily Digest from manifest.yaml')
    parser.add_argument('--manifest', default=os.environ.get('MANIFEST_PATH','manifest.yaml'))
    parser.add_argument('--out',      default=None)
    parser.add_argument('--img-dir',  default=None)
    args = parser.parse_args()

    with open(args.manifest) as f:
        meta = yaml.safe_load(f).get('report',{})

    # Allow env var overrides (useful in Docker)
    if os.environ.get('REPORT_ORG'):
        meta['organization'] = os.environ['REPORT_ORG']
    if os.environ.get('REPORT_DATE'):
        meta['date'] = os.environ['REPORT_DATE']

    if args.out is None:
        date_str = meta.get('date', datetime.now().strftime('%d %B %Y'))
        safe_date = date_str.replace(' ','_')
        org_safe  = meta.get('organization','Nokia_Networks').replace(' ','_')
        args.out = f'/app/output/{org_safe}_CTI_Digest_{safe_date}.docx'
        if not os.path.exists('/app/output'):
            args.out = f'output/{org_safe}_CTI_Digest_{safe_date}.docx'

    build_report(args.manifest, args.out, args.img_dir)
